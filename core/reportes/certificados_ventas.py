# -*- coding: utf-8 -*-
"""
Certificados de ventas mensuales (JIPER / Milagros).

Genera el certificado en Word sobre la hoja membretada correspondiente y,
si hay LibreOffice disponible, lo convierte a PDF (para firma manual).

El valor se puede tomar automaticamente del informe de ventas por centro de
costo (mismo que usa el P&G) o escribirse a mano.

Contadora unica: Luz Aida Hernandez Garcia (T.P. 159803-T, C.C. 32.297.029).
"""
from __future__ import annotations

import io
import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Cm

# ------------------------------------------------------------------ plantillas
_DIR = Path(__file__).resolve().parent / "plantillas_certificados"
PLANTILLA = {
    "JIPER": _DIR / "plantilla_jiper.docx",
    "MILAGROS": _DIR / "plantilla_milagros.docx",
}

# ------------------------------------------------------------------ contadora
CONTADORA = "LUZ AIDA HERNANDEZ GARCIA"
CC_CONTADORA = "32.297.029"
TP_CONTADORA = "159803-T"
FIRMA_IMG = _DIR / "firma_luz_aida.png"


def hay_firma() -> bool:
    return FIRMA_IMG.exists()

_MESES = ["", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio",
          "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]

# ==================================================================
#  REGISTRO DE PUNTOS DE VENTA
#  tipo: BRUTAS | VENDIO | EXITO | ARAUCO | EXITO_MIL | ARAUCO_MIL
#  cc4 : centro de costo para traer el valor del informe ("" = solo manual)
# ==================================================================
PUNTOS: List[Dict[str, str]] = [
    # ---- JIPER (membrete JIPER) ----
    {"clave": "Santa Leña Oviedo", "membrete": "JIPER", "tipo": "BRUTAS",
     "ciudad": "Rionegro", "empresa": "JIPER S.A.S", "nit": "901.038.325-1",
     "establecimiento": "SANTA LEÑA OVIEDO", "cc4": "1102"},
    {"clave": "Santa Leña El Tesoro", "membrete": "JIPER", "tipo": "BRUTAS",
     "ciudad": "Rionegro", "empresa": "JIPER S.A.S", "nit": "901.038.325-1",
     "establecimiento": "SANTA LEÑA EL TESORO", "cc4": "1103"},
    {"clave": "Santa Leña Llanogrande", "membrete": "JIPER", "tipo": "BRUTAS",
     "ciudad": "Rionegro", "empresa": "JIPER S.A.S", "nit": "901.038.325-1",
     "establecimiento": "SANTA LEÑA LLANOGRANDE", "cc4": "1108"},
    {"clave": "Santa Leña Las Vegas", "membrete": "JIPER", "tipo": "BRUTAS",
     "ciudad": "Rionegro", "empresa": "JIPER S.A.S", "nit": "901.038.325-1",
     "establecimiento": "SANTA LEÑA LAS VEGAS", "cc4": "1115"},
    {"clave": "Santa Leña Medical Tower", "membrete": "JIPER", "tipo": "BRUTAS",
     "ciudad": "Rionegro", "empresa": "JIPER S.A.S", "nit": "901.038.325-1",
     "establecimiento": "SANTA LEÑA MEDICAL TOWER", "cc4": "1116"},
    {"clave": "Santa Leña Hotel Nock", "membrete": "JIPER", "tipo": "BRUTAS",
     "ciudad": "Rionegro", "empresa": "JIPER S.A.S", "nit": "901.038.325-1",
     "establecimiento": "SANTA LEÑA HOTEL NOCK", "cc4": "1118"},
    {"clave": "Santa Leña Viva Envigado", "membrete": "JIPER", "tipo": "EXITO",
     "ciudad": "Medellín", "empresa": "JIPER S.A.S.", "nit": "901.038.325-1",
     "establecimiento": "", "marca": "Santa Leña", "centro": "VIVA ENVIGADO",
     "local": "212Z", "destinatario": "GRUPO ÉXITO S.A.", "cc4": "1106"},
    {"clave": "Santa Leña Éxito El Poblado", "membrete": "JIPER", "tipo": "EXITO",
     "ciudad": "Medellín", "empresa": "JIPER S.A.S.", "nit": "901.038.325-1",
     "establecimiento": "", "marca": "Santa Leña", "centro": "Éxito El Poblado",
     "local": "L0138", "destinatario": "GRUPO ÉXITO S.A.", "cc4": ""},
    {"clave": "Santa Leña Los Molinos", "membrete": "JIPER", "tipo": "EXITO",
     "ciudad": "Medellín", "empresa": "JIPER S.A.S.", "nit": "901.038.325-1",
     "establecimiento": "", "marca": "Santa Leña", "centro": "ALMACEN EXITO - LOS MOLINOS",
     "local": "L0112", "destinatario": "GRUPO ÉXITO S.A.", "cc4": "1113"},
    {"clave": "Santa Leña Fabricato (Parque Arauco)", "membrete": "JIPER", "tipo": "ARAUCO",
     "ciudad": "Medellín", "empresa": "JIPER S.A.S", "nit": "901.038.325-1",
     "establecimiento": "", "marca": "SANTA LEÑA", "centro": "Parque Fabricato",
     "local": "Piso 3, LC 3021", "destinatario": "PARQUE ARAUCO.", "cc4": "1112"},
    {"clave": "Milagros Manila (JIPER)", "membrete": "JIPER", "tipo": "BRUTAS",
     "ciudad": "Rionegro", "empresa": "JIPER S.A.S", "nit": "901.038.325-1",
     "establecimiento": "MILAGROS MANILA", "cc4": "1206"},
    # ---- MILAGROS (membrete MILAGROS) ----
    {"clave": "Milagros Manila", "membrete": "MILAGROS", "tipo": "VENDIO",
     "ciudad": "Medellín", "empresa": "MILAGROS GROUP S.A.S", "nit": "900.495.853-4",
     "establecimiento": "MILAGROS MANILA", "cc4": "1206"},
    {"clave": "Milagros Tesoro", "membrete": "MILAGROS", "tipo": "VENDIO",
     "ciudad": "Medellín", "empresa": "GRUPO DE ALIMENTOS DE COLOMBIA S.A.S", "nit": "900.976.364-9",
     "establecimiento": "MILAGROS TESORO", "cc4": "1201"},
    {"clave": "Milagros Viva Envigado", "membrete": "MILAGROS", "tipo": "EXITO_MIL",
     "ciudad": "Envigado", "empresa": "GRUPO DE ALIMENTOS DE COLOMBIA S.A.S", "nit": "900.976.364-9",
     "establecimiento": "MILAGROS VIVA ENVIGADO", "marca": "Milagros", "centro": "VIVA ENVIGADO",
     "local": "370", "destinatario": "GRUPO ÉXITO S.A.", "cc4": "1203"},
    {"clave": "Milagros Fabricato (Parque Arauco)", "membrete": "MILAGROS", "tipo": "ARAUCO_MIL",
     "ciudad": "Medellín", "empresa": "GRUPO DE ALIMENTOS DE COLOMBIA S.A.S", "nit": "900.976.364-9",
     "establecimiento": "MILAGROS FABRICATO", "marca": "Milagros", "centro": "Parque Fabricato",
     "local": "", "destinatario": "PARQUE ARAUCO", "cc4": "1205"},
]


def punto_por_clave(clave: str) -> Optional[Dict[str, str]]:
    for p in PUNTOS:
        if p["clave"] == clave:
            return p
    return None


# ==================================================================
#  NUMEROS A LETRAS (espanol)  -  validado contra los certificados reales
# ==================================================================
_ESP = {1: "UN", 2: "DOS", 3: "TRES", 4: "CUATRO", 5: "CINCO", 6: "SEIS",
        7: "SIETE", 8: "OCHO", 9: "NUEVE", 10: "DIEZ", 11: "ONCE", 12: "DOCE",
        13: "TRECE", 14: "CATORCE", 15: "QUINCE", 16: "DIECISEIS",
        17: "DIECISIETE", 18: "DIECIOCHO", 19: "DIECINUEVE", 20: "VEINTE",
        21: "VEINTIUN", 22: "VEINTIDOS", 23: "VEINTITRES", 24: "VEINTICUATRO",
        25: "VEINTICINCO", 26: "VEINTISEIS", 27: "VEINTISIETE", 28: "VEINTIOCHO",
        29: "VEINTINUEVE"}
_DEC = {2: "VEINTE", 3: "TREINTA", 4: "CUARENTA", 5: "CINCUENTA", 6: "SESENTA",
        7: "SETENTA", 8: "OCHENTA", 9: "NOVENTA"}
_CENT = {1: "CIENTO", 2: "DOSCIENTOS", 3: "TRESCIENTOS", 4: "CUATROCIENTOS",
         5: "QUINIENTOS", 6: "SEISCIENTOS", 7: "SETECIENTOS", 8: "OCHOCIENTOS",
         9: "NOVECIENTOS"}


def _dos(n: int) -> str:
    if n <= 29:
        return _ESP[n]
    d, u = divmod(n, 10)
    return _DEC[d] if u == 0 else _DEC[d] + " Y " + _ESP[u]


def _grupo3(n: int) -> str:
    if n == 100:
        return "CIEN"
    c, r = divmod(n, 100)
    s = (_CENT[c] + " ") if c else ""
    if r:
        s += _dos(r)
    return s.strip()


def _seccion6(n: int) -> str:
    miles, uni = divmod(n, 1000)
    s = ("MIL " if miles == 1 else _grupo3(miles) + " MIL ") if miles else ""
    if uni:
        s += _grupo3(uni)
    return s.strip()


def entero_en_letras(n: int) -> str:
    """Entero (0..billones) en MAYUSCULAS, sin la palabra 'pesos'."""
    n = int(n)
    if n == 0:
        return "CERO"
    s = ""
    billones, n = divmod(n, 10 ** 12)
    millones, n = divmod(n, 10 ** 6)
    uni = n
    if billones:
        s += "UN BILLON " if billones == 1 else _seccion6(billones) + " BILLONES "
    if millones:
        s += "UN MILLON " if millones == 1 else _seccion6(millones) + " MILLONES "
    if uni:
        s += _seccion6(uni)
    return s.strip()


def _titulo(s: str) -> str:
    """MAYUSCULAS -> Formato Titulo, con la conjuncion 'y' en minuscula."""
    t = s.lower().title()
    return (" " + t + " ").replace(" Y ", " y ").strip()


def peso_o_pesos(n: int) -> str:
    n = int(n)
    return "Peso" if (n % 10 == 1 and n % 100 != 11) else "Pesos"


def formato_pesos(n: float) -> str:
    return format(int(round(n)), ",d").replace(",", ".")


# ==================================================================
#  CONSTRUCCION DEL DOCUMENTO
# ==================================================================
def _par(doc, texto="", *, bold=False, center=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if texto != "":
        r = p.add_run(texto)
        r.font.name = "Arial"
        r.font.size = Pt(size)
        r.bold = bold
    return p


def _run(p, texto, *, bold=False, size=12):
    r = p.add_run(texto)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.bold = bold
    return r


def _cuerpo_brutas(doc, pt, mes, val, letras_up, peso):
    _par(doc, "CERTIFICADO", bold=True, center=True)
    _par(doc); _par(doc)
    p = _par(doc)
    _run(p, "Yo, ")
    _run(p, pt["contador"], bold=True)
    _run(p, ", identificada con CÉDULA DE CIUDADANÍA " + pt["cc"] +
         ", en calidad de Contadora de la empresa ")
    _run(p, pt["empresa"], bold=True)
    _run(p, ", identificada con NIT " + pt["nit"] +
         ". Certifico que el establecimiento de comercio ")
    _run(p, pt["establecimiento"], bold=True)
    _run(p, ", tuvo Ventas brutas en el mes de " + mes + " por valor de: ")
    _run(p, val + " ML.", bold=True)
    _run(p, " (" + letras_up + " " + peso.upper() + " ML)")
    _par(doc)
    _par(doc, "Cualquier inquietud con gusto será atendida.")
    _par(doc)
    _par(doc, "Cordialmente,")


def _cuerpo_vendio(doc, pt, mes, val, letras_ti, peso):
    _par(doc, "ASUNTO: CERTIFICACION DE VENTAS MENSUAL", bold=True)
    _par(doc)
    p = _par(doc)
    _run(p, "Yo, ")
    _run(p, pt["contador"], bold=True)
    _run(p, " identificada con cédula de ciudadanía número " + pt["cc"] +
         ", en calidad de Contadora de la empresa ")
    _run(p, pt["empresa"], bold=True)
    _run(p, ", identificada con el número de NIT: " + pt["nit"] +
         " Certifico que el establecimiento de comercio ")
    _run(p, pt["establecimiento"], bold=True)
    _run(p, ", vendió durante el mes de " + mes + ", ")
    _run(p, "$" + val, bold=True)
    _run(p, " (" + letras_ti + " " + peso + " ML).")
    _par(doc)
    _par(doc, "Cualquier inquietud con gusto será atendida.")
    _par(doc)
    _par(doc, "Cordialmente,")


def _tabla_mes_valor(doc, mes, val, etiqueta, encabezado):
    p = _par(doc, encabezado, bold=True)
    p2 = _par(doc)
    _run(p2, mes)
    _run(p2, "\t$" + val)


def _cuerpo_exito(doc, pt, mes, val):
    _par(doc, "Señores; " + pt["destinatario"])
    _par(doc, "Vicepresidencia Inmobiliaria")
    _par(doc)
    _par(doc, "ASUNTO: CERTIFICACION DE VENTAS MENSUAL", bold=True)
    _par(doc)
    _par(doc, "Apreciados señores en atención al compromiso contractual, por la "
              "presente confirmo información requerida así:")
    _par(doc)
    _par(doc, "RAZON SOCIAL: " + pt["empresa"])
    _par(doc, "NIT: " + pt["nit"])
    _par(doc, "MARCA: " + pt.get("marca", ""))
    _par(doc, "CENTRO COMERCIAL: " + pt.get("centro", ""))
    _par(doc, "# LOCAL: " + pt.get("local", ""))
    _par(doc)
    _tabla_mes_valor(doc, mes, val, "IPO", "MES – AÑO\tVENTA ANTES DE IVA/IPO")
    _par(doc)
    _par(doc, "Atentamente,")


def _cuerpo_arauco(doc, pt, mes, val):
    _par(doc, "CERTIFICACIÓN DE VENTAS", bold=True, center=True)
    _par(doc)
    _par(doc, "Señores")
    _par(doc, pt["destinatario"])
    _par(doc)
    _par(doc, "Apreciados señores, en mi calidad de Contadora y en atención al "
              "compromiso contractual, por la presente confirmo la información "
              "requerida así:")
    _par(doc)
    _par(doc, "RAZON SOCIAL: " + pt["empresa"])
    _par(doc, "NIT: " + pt["nit"])
    _par(doc, "MARCA: " + pt.get("marca", ""))
    _par(doc, "CENTRO COMERCIAL: " + pt.get("centro", ""))
    _par(doc, "# LOCAL: " + pt.get("local", ""))
    _par(doc)
    _tabla_mes_valor(doc, mes, val, "IVA", "Mes – Año\tVENTA ANTES DE IVA")
    _par(doc)
    _par(doc, "Cualquier inquietud adicional con gusto la atenderemos en los "
              "teléfonos o dirección preimpresos.")
    _par(doc)
    _par(doc, "Cordialmente,")


def _cuerpo_exito_mil(doc, pt, mes, val):
    _par(doc, "ASUNTO: CERTIFICACION DE VENTAS MENSUAL", bold=True)
    _par(doc)
    p = _par(doc)
    _run(p, "Certifico que las ventas obtenidas por la compañía ")
    _run(p, pt["empresa"], bold=True)
    _run(p, " Identificada con el número de NIT: " + pt["nit"] +
         ", en el punto de venta " + pt.get("centro", "") +
         ", LOCAL " + pt.get("local", "") + " (Marca " + pt.get("marca", "") +
         "), fueron por valor de:")
    _par(doc)
    _par(doc, "Señores;")
    _par(doc, pt["destinatario"])
    _par(doc, "Vicepresidencia Inmobiliaria")
    _par(doc)
    _tabla_mes_valor(doc, mes, val, "IPO", "MES – AÑO\tVENTA ANTES DE IVA/IPO")
    _par(doc)
    _par(doc, "Cordialmente,")


def _cuerpo_arauco_mil(doc, pt, mes, val):
    _par(doc, "ASUNTO: CERTIFICACION DE VENTAS MENSUAL", bold=True)
    _par(doc)
    p = _par(doc)
    _run(p, "Yo, ")
    _run(p, pt["contador"], bold=True)
    _run(p, " identificada con cédula de ciudadanía " + pt["cc"] +
         ", en calidad de Contadora de la empresa ")
    _run(p, pt["empresa"], bold=True)
    _run(p, ", identificada con el número de NIT: " + pt["nit"] +
         " Certifico que el establecimiento de comercio ")
    _run(p, pt["establecimiento"], bold=True)
    _run(p, ", vendió durante el mes:")
    _par(doc)
    _par(doc, "Señores;")
    _par(doc, pt["destinatario"])
    _par(doc)
    _tabla_mes_valor(doc, mes, val, "IPO", "MES – AÑO\tVENTA ANTES DE IVA/IPO")
    _par(doc)
    _par(doc, "Cordialmente,")


def _bloque_firma(doc, firmar=False):
    firmado = firmar and FIRMA_IMG.exists()
    if firmado:
        _par(doc)
        # Firma (PNG sin fondo) apoyada justo sobre la linea, que a su vez queda
        # encima del nombre. Espaciado ajustado para que no quede muy arriba.
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        try:
            p.add_run().add_picture(str(FIRMA_IMG), width=Cm(4.2))
        except Exception:
            _par(doc)
        pl = doc.add_paragraph("__________________________________")
        pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pl.paragraph_format.space_before = Pt(0)
        pl.paragraph_format.space_after = Pt(0)
        for r in pl.runs:
            r.font.name = "Arial"; r.font.size = Pt(12)
    else:
        _par(doc)
        _par(doc)
        _par(doc, "__________________________________")
    _par(doc, CONTADORA, bold=True)
    _par(doc, "Contadora Pública")
    _par(doc, "T.P. " + TP_CONTADORA, bold=True)


def generar_certificado_docx(clave: str, mes_texto: str, valor: float,
                             firmar: bool = False) -> bytes:
    """Devuelve el certificado en Word (.docx) como bytes.

    firmar=True inserta la firma escaneada de la contadora sobre la linea.
    """
    pt = dict(punto_por_clave(clave) or {})
    if not pt:
        raise ValueError("Punto de venta no encontrado: %s" % clave)
    # datos de la contadora unica
    pt["contador"] = CONTADORA
    pt["cc"] = CC_CONTADORA
    pt["tp"] = TP_CONTADORA

    mes = (mes_texto or "").strip().upper()
    val = formato_pesos(valor)
    up = entero_en_letras(valor)
    ti = _titulo(up)
    peso = peso_o_pesos(valor)

    tmpl = PLANTILLA[pt["membrete"]]
    doc = Document(str(tmpl))
    # la plantilla trae un parrafo vacio; lo usamos como primer parrafo del cuerpo
    # (borramos su contenido, agregamos parrafos nuevos despues)

    # fecha: solo en membrete JIPER (Milagros trae fecha automatica en el encabezado)
    if pt["membrete"] == "JIPER":
        from datetime import date
        hoy = date.today()
        fecha = "%s, %02d de %s de %d" % (pt["ciudad"], hoy.day, _MESES[hoy.month], hoy.year)
        _par(doc, fecha, bold=True)
        _par(doc)

    tipo = pt["tipo"]
    if tipo == "BRUTAS":
        _cuerpo_brutas(doc, pt, mes, val, up, peso)
    elif tipo == "VENDIO":
        _cuerpo_vendio(doc, pt, mes, val, ti, peso)
    elif tipo == "EXITO":
        _cuerpo_exito(doc, pt, mes, val)
    elif tipo == "ARAUCO":
        _cuerpo_arauco(doc, pt, mes, val)
    elif tipo == "EXITO_MIL":
        _cuerpo_exito_mil(doc, pt, mes, val)
    elif tipo == "ARAUCO_MIL":
        _cuerpo_arauco_mil(doc, pt, mes, val)
    else:
        _par(doc, "CERTIFICADO", bold=True, center=True)

    _bloque_firma(doc, firmar=firmar)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ==================================================================
#  CONVERSION A PDF (LibreOffice)
# ==================================================================
def _soffice_bin() -> Optional[str]:
    for name in ("libreoffice", "soffice"):
        path = shutil.which(name)
        if path:
            return path
    return None


def hay_pdf() -> bool:
    """True si el servidor puede convertir a PDF."""
    return _soffice_bin() is not None


def docx_a_pdf(docx_bytes: bytes) -> bytes:
    """Convierte un .docx (bytes) a PDF (bytes) usando LibreOffice."""
    binp = _soffice_bin()
    if not binp:
        raise RuntimeError("LibreOffice no esta instalado en el servidor; "
                           "no es posible generar PDF.")
    with tempfile.TemporaryDirectory() as d:
        src = os.path.join(d, "cert.docx")
        with open(src, "wb") as f:
            f.write(docx_bytes)
        subprocess.run([binp, "--headless", "--convert-to", "pdf",
                        "--outdir", d, src],
                       check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                       timeout=120)
        pdf = os.path.join(d, "cert.pdf")
        with open(pdf, "rb") as f:
            return f.read()


def generar_certificado_pdf(clave: str, mes_texto: str, valor: float,
                            firmar: bool = False) -> bytes:
    return docx_a_pdf(generar_certificado_docx(clave, mes_texto, valor, firmar=firmar))
